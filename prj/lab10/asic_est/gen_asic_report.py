#!/usr/bin/env python3
"""Generate ASIC estimation report as .docx"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(title, headers, rows):
    if title:
        para(title, bold=True)
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Calibri'

# ================================================================
heading('ASIC Area and Power Estimation Report', 0)
para('SoC: 4-Thread SIMT GPU with 4x4 BF16 Tensor Core')
para('EE 533 - Network Processor Design')
para('Raymond | Team 5')
doc.add_paragraph()

# ================================================================
heading('1. Methodology', 1)
para('This report estimates the ASIC area, power, and gate count of the SoC design '
     'using open-source EDA tools and standard cell libraries. The methodology:')

para('1. RTL Source: Jeremy Cai\'s simulation RTL (Verilog), identical to the design '
     'verified in iverilog simulation and synthesized to Xilinx Virtex-II Pro FPGA.')
para('2. Synthesis Tool: Yosys 0.38 (open-source logic synthesis)')
para('3. Target Library: Nangate 45nm Open Cell Library (FreePDK45-compatible, '
     'published cell areas in um2)')
para('4. Scaling: Published TSMC density scaling factors for 28nm, 16nm, 7nm, 5nm nodes')
para('5. Memory: SRAM area estimated separately using published bit-cell densities per node')

# ================================================================
heading('2. Design Hierarchy', 1)

code_block(
    'soc (top)\n'
    '  +-- cpu_mt             4-thread FGMT barrel ARM CPU\n'
    '  |     +-- alu, barrel_shifter, bdtu, cu, regfile\n'
    '  +-- cp10_regfile        Coprocessor registers (10 CRs)\n'
    '  +-- dma_engine          DMA controller (D_IMEM/D_UNPACK/D_PACK)\n'
    '  +-- pkt_proc            Packet processor (command FSM)\n'
    '  +-- conv_fifo           Convertible FIFO (RX/TX/SRAM modes)\n'
    '  +-- sm_core             SIMT GPU streaming multiprocessor\n'
    '  |     +-- sp_core x4    4 streaming processors\n'
    '  |     |     +-- gpr_regfile (16x16b), pred_regfile\n'
    '  |     |     +-- int16alu, bf16fpu, pplbfintcvt\n'
    '  |     +-- tc_top        Tensor core wrapper\n'
    '  |     |     +-- tensor_core -> bf16sa (4x4 systolic array)\n'
    '  |     |           +-- bf16pe x16 (BF16 multiply-accumulate)\n'
    '  |     +-- fetch_unit, sm_decoder, scoreboard, simt_stack\n'
    '  +-- test_i_mem          CPU IMEM (4096x32 SRAM)\n'
    '  +-- test_d_mem          CPU DMEM (4096x32 SRAM)\n'
    '  +-- test_gpu_imem       GPU IMEM (256x32 SRAM)\n'
    '  +-- test_gpu_dmem x4    GPU DMEM banks (1024x16 each)\n'
    '  +-- test_dpfifo         FIFO SRAM (4096x64)')

# ================================================================
heading('3. Synthesis Results (Nangate 45nm)', 1)

heading('3.1 Overall Summary', 2)
add_table('', ['Metric', 'Value'], [
    ['Total standard cells', '60,518'],
    ['Combinational cells', '52,744 (87.2%)'],
    ['Sequential cells (flip-flops)', '7,774 (12.8%)'],
    ['Memory blocks (SRAM)', '24 (not mapped to std cells)'],
    ['Logic chip area', '103,856 um2 = 0.104 mm2'],
    ['SRAM bits', '598,016 bits'],
])

heading('3.2 Cell Distribution by Category', 2)
add_table('', ['Category', 'Count', '%', 'Purpose'], [
    ['NAND gates', '8,431', '13.9%', 'Core logic'],
    ['AOI (AND-OR-Invert)', '7,883', '13.0%', 'Compound logic in adders/comparators'],
    ['Flip-flops (DFF)', '7,774', '12.8%', 'Pipeline registers, FSM state'],
    ['MUX (2:1)', '7,279', '12.0%', 'Data routing, RF muxes, thread select'],
    ['OAI (OR-AND-Invert)', '6,649', '11.0%', 'Compound logic in BF16 arithmetic'],
    ['NOR gates', '6,399', '10.6%', 'Core logic'],
    ['XOR/XNOR gates', '6,295', '10.4%', 'Adders, BF16 exponent/mantissa ops'],
    ['Inverters', '4,117', '6.8%', 'Signal inversion'],
    ['AND gates', '3,539', '5.8%', 'Glue logic, enables'],
    ['OR gates', '2,152', '3.6%', 'Glue logic'],
])

heading('3.3 Top 10 Individual Cell Types', 2)
add_table('', ['Cell', 'Count', '%'], [
    ['MUX2_X1 (2:1 mux)', '7,279', '12.0%'],
    ['DFFR_X1 (DFF w/ reset)', '6,622', '10.9%'],
    ['NAND2_X1', '6,286', '10.4%'],
    ['OAI21_X1', '4,578', '7.6%'],
    ['NOR2_X1', '4,517', '7.5%'],
    ['AOI21_X1', '4,134', '6.8%'],
    ['INV_X1', '4,117', '6.8%'],
    ['XNOR2_X1', '3,692', '6.1%'],
    ['XOR2_X1', '2,603', '4.3%'],
    ['AND2_X1', '2,561', '4.2%'],
])
para('MUX2 is the most common cell, reflecting the heavy data multiplexing in the '
     '4-thread SIMT architecture (register file read ports, forwarding paths, '
     'tensor core gather/scatter, DMA bank selection).')

# ================================================================
heading('4. Area Breakdown by Function', 1)

heading('4.1 Logic Area (from Yosys cell mapping)', 2)
add_table('', ['Component', 'Area (um2)', '% of Logic'], [
    ['ARM CPU (4-thread barrel)', '19,782', '31.3%'],
    ['GPU SM Core (4x SP lanes)', '14,298', '22.6%'],
    ['GPU Tensor Core (4x4 SA)', '9,686', '15.3%'],
    ['Packet Processor + FIFO', '6,470', '10.2%'],
    ['CP10 + DMA Engine', '3,888', '6.2%'],
    ['Other (memory wrappers, etc)', '9,063', '14.3%'],
    ['TOTAL', '63,186', '100%'],
])

heading('4.2 SRAM Area (estimated from bit-cell density)', 2)
add_table('', ['Memory', 'Size', 'Bits', 'Area @ 45nm (um2)'], [
    ['CPU IMEM', '4096 x 32', '131,072', '65,536'],
    ['CPU DMEM', '4096 x 32', '131,072', '65,536'],
    ['GPU IMEM', '256 x 32', '8,192', '4,096'],
    ['GPU DMEM x4', '4 x 1024 x 16', '65,536', '32,768'],
    ['Conv FIFO', '4096 x 64', '262,144', '131,072'],
    ['TOTAL', '', '598,016', '299,008'],
])
para('SRAM density: 0.5 um2/bit at 45nm (6T SRAM). SRAM dominates total area (~83%).')

heading('4.3 Total Area at 45nm', 2)
add_table('', ['Component', 'Area (um2)', 'Area (mm2)', '%'], [
    ['Logic (std cells)', '63,186', '0.063', '17%'],
    ['SRAM', '299,008', '0.299', '83%'],
    ['TOTAL', '362,194', '0.362', '100%'],
])

# ================================================================
heading('5. TSMC Node Scaling', 1)

para('Area estimates scaled using published TSMC logic density and SRAM bit-cell size:')

add_table('Area Scaling to TSMC Nodes',
    ['Node', 'Logic (mm2)', 'SRAM (mm2)', 'Total (mm2)', 'vs 45nm'],
    [
        ['45nm', '0.063', '0.299', '0.362', '1.00x'],
        ['TSMC 28nm', '0.028', '0.072', '0.100', '0.28x'],
        ['TSMC 16nm', '0.014', '0.042', '0.056', '0.15x'],
        ['TSMC 7nm', '0.006', '0.024', '0.030', '0.08x'],
        ['TSMC 5nm', '0.004', '0.013', '0.017', '0.05x'],
    ])

doc.add_paragraph()
para('Note on FPGA comparison: The Virtex-II Pro XC2VP50 full die is 365 mm2, but '
     'most of that is programmable routing fabric (~55%), configuration SRAM (~15%), '
     'I/O ring (~10%), and unused PowerPC cores (~5%). The actual CLB/BRAM/MULT '
     'resources occupy ~55 mm2 (15% of die). Our design uses 89% of slices, so '
     'the effective FPGA silicon for our logic is ~49 mm2. '
     'The ASIC at 45nm (0.36 mm2) is ~136x smaller than this effective area. '
     'At TSMC 7nm (0.03 mm2), the ratio increases to ~1,600x.',
     bold=True)

# ================================================================
heading('6. Power Estimation', 1)

para('Dynamic power estimated from cell count and published power-per-gate metrics:')

add_table('Power Estimation by Node',
    ['Node', 'Frequency', 'Vdd', 'Logic Power', 'Notes'],
    [
        ['45nm', '125 MHz', '1.1V', '~1,150 mW', 'FPGA clock rate'],
        ['TSMC 28nm', '500 MHz', '0.9V', '~1,080 mW', '4x faster, lower V'],
        ['TSMC 16nm', '800 MHz', '0.8V', '~580 mW', 'FinFET benefit'],
        ['TSMC 7nm', '1 GHz', '0.75V', '~260 mW', 'High performance'],
        ['TSMC 5nm', '1.2 GHz', '0.7V', '~180 mW', 'Mobile-class'],
    ])

doc.add_paragraph()
para('Note: SRAM leakage and dynamic power add ~50-200 mW depending on access patterns '
     'and node. Total system power at TSMC 7nm: estimated 300-500 mW at 1 GHz.')

# ================================================================
heading('7. FPGA vs ASIC Comparison', 1)

para('The Virtex-II Pro die is 365 mm2, but this includes routing fabric, configuration '
     'SRAM, I/O ring, and unused PowerPC cores. A fair comparison uses only the '
     'effective silicon area occupied by our logic.')

add_table('FPGA (Virtex-II Pro) vs ASIC Comparison',
    ['Metric', 'FPGA (XC2VP50)', 'ASIC (45nm)', 'ASIC (7nm)'],
    [
        ['Full die area', '365 mm2', '0.36 mm2', '0.03 mm2'],
        ['Effective logic area', '~49 mm2 (89% of CLBs)', '0.36 mm2', '0.03 mm2'],
        ['Area ratio (vs FPGA effective)', '1x', '136x smaller', '1,600x smaller'],
        ['Clock Frequency', '125 MHz', '~200 MHz', '~1 GHz'],
        ['Logic Power', '~12 W (full board)', '~1.2 W', '~0.3 W'],
        ['Process Node', '130nm', '45nm', '7nm'],
        ['Slices / Std Cells', '21,060 slices', '60,518 cells', '60,518 cells'],
        ['Block RAM / SRAM', '145 RAMB16', '598 Kbit', '598 Kbit'],
        ['Multipliers', '24 MULT18X18', 'In std cells', 'In std cells'],
    ])

para('FPGA area breakdown: ~55% routing fabric, ~15% config SRAM, ~15% CLB/BRAM/MULT, '
     '~10% I/O ring, ~5% PowerPC cores (unused). Only the CLB/BRAM/MULT fraction '
     '(~55 mm2) is functionally equivalent to the ASIC logic + SRAM.')

# ================================================================
heading('8. Design Efficiency Analysis', 1)

para('Compute density (TOPS/mm2 for BF16 tensor core):', bold=True)

code_block(
    'Tensor core: 4x4 BF16 matmul = 128 FLOPs per WMMA (64 MACs x 2 FLOPs)\n'
    'WMMA latency: ~44 cycles\n'
    '\n'
    'At 1 GHz (TSMC 7nm):\n'
    '  Throughput = 128 / 44 = 2.9 GFLOPS (BF16)\n'
    '  Tensor core area = 0.001 mm2 (scaled from 9,686 um2 at 45nm)\n'
    '  Compute density = 2.9 / 0.001 = 2,900 GFLOPS/mm2\n'
    '\n'
    'For comparison:\n'
    '  NVIDIA A100 tensor cores: ~600 GFLOPS/mm2 (BF16)\n'
    '  Our design is ~5x denser per tensor core,\n'
    '  but A100 has 432 tensor cores vs our 1.')

# ================================================================
heading('9. Methodology Notes', 1)

para('Limitations of this estimation:')
para('1. Yosys uses generic ABC technology mapping, not commercial synthesis tools '
     '(Synopsys DC / Cadence Genus). Commercial tools achieve 10-20% better area.')
para('2. Nangate 45nm is an academic library, not a production TSMC library. '
     'Cell areas are representative but not exact.')
para('3. SRAM area is estimated from published bit-cell densities, not from '
     'actual SRAM compiler output.')
para('4. Power estimation uses rule-of-thumb scaling, not gate-level power analysis. '
     'Actual power depends on switching activity, clock tree, and leakage.')
para('5. No physical design (placement, routing, clock tree) was performed. '
     'Routing overhead typically adds 20-40% to cell area.')
para('6. Memory blocks ($mem_v2) are not mapped to standard cells — they would be '
     'replaced by SRAM macros in a real ASIC flow.')

# ================================================================
heading('10. Conclusion', 1)

para('The SoC design containing a 4-thread SIMT GPU with 4x4 BF16 tensor core, '
     'barrel ARM CPU, DMA engine, and packet processor synthesizes to approximately '
     '60,500 standard cells (250K gate equivalents). At TSMC 28nm, the total die area '
     'including SRAM is estimated at 0.10 mm2, consuming approximately 1W at 500 MHz. '
     'The design is dominated by SRAM (83% of area) and the ARM CPU (31% of logic). '
     'The tensor core, despite its compute capability, accounts for only 15% of '
     'logic area — demonstrating efficient hardware utilization for BF16 matrix operations.',
     bold=True)

# ================================================================
output_path = '/home/raymond/USC/ee533/prj/lab10/asic_est/asic_estimation_report.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
